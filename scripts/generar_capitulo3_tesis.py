from __future__ import annotations

from collections import Counter
from io import BytesIO
from pathlib import Path
from zipfile import ZipFile
import json
import os
import urllib.request

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(r"D:\Proyecto_FCI_2025")
DOWNLOADS = Path(r"C:\Users\R-R\Downloads")
ASSETS_DIR = ROOT / "tmp" / "chapter3-assets"
OUTPUT_DIR = ROOT / "output" / "docx"
ORIGINAL_DOC = DOWNLOADS / "PLANTILLA PARA TRABAJO DE INTEGRACIÓN CURRICULAR.docx"
OUTPUT_DOC = OUTPUT_DIR / "PLANTILLA_TRABAJO_INTEGRACION_CURRICULAR_capitulo3.docx"
DOWNLOAD_COPY = DOWNLOADS / "PLANTILLA PARA TRABAJO DE INTEGRACIÓN CURRICULAR - Capítulo 3 integrado.docx"
SCHEMA_IMG = DOWNLOADS / "supabase-schema-mvfmxecpltbhvjcneisf.png"


def load_env(path: Path) -> dict[str, str]:
    values: dict[str, str] = {}
    if not path.exists():
        return values
    for raw_line in path.read_text(encoding="utf-8").splitlines():
        line = raw_line.strip()
        if not line or line.startswith("#") or "=" not in line:
            continue
        key, value = line.split("=", 1)
        values[key.strip()] = value.strip().strip('"').strip("'")
    return values


def fetch_questionnaire_data() -> list[dict]:
    env = load_env(ROOT / ".env.local")
    url = env.get("NEXT_PUBLIC_SUPABASE_URL")
    key = env.get("SUPABASE_SERVICE_ROLE_KEY")
    if not url or not key:
        return []

    endpoint = (
        f"{url}/rest/v1/cuestionario_limpio_respuestas"
        "?select=parroquia,sector_economico,modalidad_preferida,interes_programa"
    )
    request = urllib.request.Request(
        endpoint,
        headers={
            "apikey": key,
            "Authorization": f"Bearer {key}",
            "Accept": "application/json",
        },
    )
    with urllib.request.urlopen(request, timeout=30) as response:
        return json.loads(response.read().decode("utf-8"))


def make_counter(rows: list[dict], field: str) -> Counter:
    values = []
    for row in rows:
        value = row.get(field)
        if value is None:
            continue
        text = str(value).strip()
        if text:
            values.append(text)
    return Counter(values)


def title_case_first(text: str) -> str:
    if not text:
        return text
    return text[:1].upper() + text[1:]


def create_architecture_diagram(output_path: Path) -> None:
    output_path.parent.mkdir(parents=True, exist_ok=True)
    width, height = 1800, 1100
    img = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(img)

    font_dir = Path(r"C:\Windows\Fonts")
    regular_path = font_dir / "arial.ttf"
    bold_path = font_dir / "arialbd.ttf"
    regular = ImageFont.truetype(str(regular_path), 24)
    regular_small = ImageFont.truetype(str(regular_path), 18)
    bold = ImageFont.truetype(str(bold_path), 28)
    title = ImageFont.truetype(str(bold_path), 42)
    subtitle = ImageFont.truetype(str(regular_path), 23)

    blue = "#0f4c97"
    blue_mid = "#1f6fbe"
    cyan = "#2ea8df"
    border = "#91b7dc"
    panel = "#eef5fc"
    panel_alt = "#f7fbff"
    text = "#18324a"
    muted = "#516679"

    draw.rounded_rectangle((60, 40, 1740, 150), radius=24, fill=blue)
    draw.text((96, 68), "Arquitectura funcional de la plataforma web propuesta", font=title, fill="white")
    draw.text(
        (96, 114),
        "Estructura por capas, servicios y módulos operativos articulados con Supabase.",
        font=subtitle,
        fill="#dcecff",
    )

    top_boxes = [
        (
            (120, 210, 1680, 360),
            "Capa de presentación",
            "Interfaz web responsiva desarrollada en Next.js para administradora, formadora, investigadora y mujer emprendedora. "
            "Incluye dashboard, formularios, reportes, seguimiento de cursos y navegación diferenciada por rol.",
        ),
        (
            (120, 405, 1680, 610),
            "Capa de aplicación y lógica de negocio",
            "Gestiona autenticación, permisos, reglas condicionales de la encuesta, publicación de cursos, actividades del proyecto, "
            "producción científica, notificaciones, historial de ingresos y generación de reportes. Se apoya en Server Actions, rutas API y validaciones.",
        ),
    ]
    for rect, head, body in top_boxes:
        draw.rounded_rectangle(rect, radius=26, fill=panel, outline=border, width=3)
        draw.text((rect[0] + 28, rect[1] + 22), head, font=bold, fill=blue)
        draw.multiline_text((rect[0] + 28, rect[1] + 70), body, font=regular, fill=text, spacing=10)

    columns = [
        (
            (120, 660, 610, 1010),
            "Servicios Supabase",
            [
                "Auth para inicio de sesión, invitaciones y recuperación de contraseña.",
                "PostgreSQL para tablas normalizadas de usuarios, encuesta, cursos, documentos y producción científica.",
                "Storage para documentos institucionales y entregas en PDF.",
                "Persistencia centralizada y control de acceso sobre la información del proyecto.",
            ],
        ),
        (
            (655, 660, 1145, 1010),
            "Módulos operativos",
            [
                "Diagnóstico y captura de respuestas del cuestionario.",
                "Analítica de necesidades y predicción de cursos.",
                "Diseño de cursos, tareas y malla formativa.",
                "Avance del proyecto, producción científica y reportes.",
            ],
        ),
        (
            (1190, 660, 1680, 1010),
            "Salidas e indicadores",
            [
                "Tarjetas del dashboard y métricas de validación.",
                "Gráficos de parroquias, sectores, interés y modalidad.",
                "Próximas actividades, alertas y notificaciones.",
                "Reportes académicos y documentos descargables.",
            ],
        ),
    ]

    for rect, head, bullets in columns:
        draw.rounded_rectangle(rect, radius=24, fill=panel_alt, outline=border, width=3)
        draw.text((rect[0] + 24, rect[1] + 22), head, font=bold, fill=blue_mid)
        y = rect[1] + 76
        for bullet in bullets:
            draw.ellipse((rect[0] + 28, y + 9, rect[0] + 42, y + 23), fill=cyan)
            draw.multiline_text((rect[0] + 58, y), bullet, font=regular, fill=text, spacing=8)
            y += 78

    for start_y, end_y in [(360, 405), (610, 660)]:
        x = width // 2
        draw.line((x, start_y, x, end_y), fill=blue_mid, width=8)
        draw.polygon([(x - 16, end_y - 20), (x + 16, end_y - 20), (x, end_y + 8)], fill=blue_mid)

    draw.text(
        (72, 1042),
        "Nota. Elaboración propia a partir de la arquitectura implementada en la plataforma y de su integración funcional con Supabase.",
        font=regular_small,
        fill=muted,
    )
    img.save(output_path)


def set_cell_text(cell, text: str, bold_value: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(text)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(11)
    run.bold = bold_value


def set_run_font(run, size: int = 12, bold: bool = False, italic: bool = False) -> None:
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def style_paragraph(paragraph, align=WD_ALIGN_PARAGRAPH.JUSTIFY, first_line: float = 0.5) -> None:
    paragraph.alignment = align
    paragraph.paragraph_format.line_spacing = 2
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.space_before = Pt(0)
    if first_line:
        paragraph.paragraph_format.first_line_indent = Inches(first_line)
    else:
        paragraph.paragraph_format.first_line_indent = Inches(0)


def add_text(document: Document, text: str, *, italic=False) -> None:
    paragraph = document.add_paragraph()
    style_paragraph(paragraph)
    run = paragraph.add_run(text)
    set_run_font(run, size=12, italic=italic)


def add_heading(document: Document, text: str, level: int) -> None:
    paragraph = document.add_paragraph()
    style_paragraph(paragraph, align=WD_ALIGN_PARAGRAPH.LEFT, first_line=0)
    run = paragraph.add_run(text)
    sizes = {1: 14, 2: 13, 3: 12}
    set_run_font(run, size=sizes.get(level, 12), bold=True)


def add_figure(document: Document, number: int, title: str, image_path: Path, note: str, width: float = 6.4) -> None:
    p1 = document.add_paragraph()
    style_paragraph(p1, align=WD_ALIGN_PARAGRAPH.LEFT, first_line=0)
    r1 = p1.add_run(f"Figura {number}")
    set_run_font(r1, size=12, bold=True)

    p2 = document.add_paragraph()
    style_paragraph(p2, align=WD_ALIGN_PARAGRAPH.LEFT, first_line=0)
    r2 = p2.add_run(title)
    set_run_font(r2, size=12, italic=True)

    p_img = document.add_paragraph()
    style_paragraph(p_img, align=WD_ALIGN_PARAGRAPH.CENTER, first_line=0)
    p_img.add_run().add_picture(str(image_path), width=Inches(width))

    p3 = document.add_paragraph()
    style_paragraph(p3, align=WD_ALIGN_PARAGRAPH.LEFT, first_line=0)
    r3 = p3.add_run(f"Nota. {note}")
    set_run_font(r3, size=10)


def add_module_subsection(document: Document, title: str, body: str) -> None:
    add_heading(document, title, 3)
    add_text(document, body)


def add_page_break(document: Document) -> None:
    paragraph = document.add_paragraph()
    paragraph.add_run().add_break(WD_BREAK.PAGE)


def main() -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    ASSETS_DIR.mkdir(parents=True, exist_ok=True)

    rows = fetch_questionnaire_data()
    total = len(rows)
    parroquias = make_counter(rows, "parroquia")
    sectores = make_counter(rows, "sector_economico")
    modalidades = make_counter(rows, "modalidad_preferida")
    intereses = make_counter(rows, "interes_programa")

    top_parroquias = parroquias.most_common(4)
    top_sectores = sectores.most_common(4)
    modalidades_text = ", ".join(f"{k}: {v}" for k, v in modalidades.most_common())
    interes_text = ", ".join(f"{k}: {v}" for k, v in intereses.most_common())

    professional_arch = ASSETS_DIR / "diagrama_arquitectura_profesional.png"
    create_architecture_diagram(professional_arch)

    document = Document(ORIGINAL_DOC)
    add_page_break(document)

    add_heading(document, "CAPÍTULO III", 1)
    add_heading(document, "DISEÑO DE LA PROPUESTA TECNOLÓGICA", 1)

    add_heading(document, "3.1. Fundamentación de la propuesta", 2)
    add_text(
        document,
        "La propuesta tecnológica desarrollada para el proyecto Kawsay Emprende Guayaquil se concibe como un sistema web integrador "
        "orientado a centralizar la gestión operativa, académica y analítica del programa. Su propósito no se limita al registro de "
        "información, sino que articula la administración de usuarios, la captura y análisis de encuestas diagnósticas, la proyección "
        "de cursos formativos, la gestión documental, el seguimiento del avance del proyecto y la producción científica en un solo entorno digital."
    )
    add_text(
        document,
        "Desde el punto de vista metodológico, la plataforma responde a un enfoque de desarrollo incremental y modular, donde cada componente "
        "se implementa como una unidad funcional capaz de interactuar con las demás sin perder independencia operativa. Este criterio favorece "
        "la mantenibilidad del sistema, la escalabilidad de nuevas funciones y la trazabilidad de la información, principios consistentes con "
        "las recomendaciones de Pressman y Maxim (2020) y Sommerville (2019) para soluciones de ingeniería de software aplicadas a contextos organizacionales."
    )

    add_heading(document, "3.1.1. Arquitectura funcional de la plataforma", 3)
    add_text(
        document,
        "La arquitectura funcional de la plataforma se estructura en tres niveles principales: una capa de presentación, una capa de "
        "aplicación y una capa de servicios y persistencia de datos. La capa de presentación corresponde a la interfaz web responsiva "
        "desarrollada en Next.js, desde la cual interactúan la administradora, la formadora, la investigadora y la mujer emprendedora. "
        "La capa de aplicación concentra la lógica de negocio relacionada con autenticación, permisos, flujos de encuesta, diseño de cursos, "
        "seguimiento de actividades, reportes y notificaciones. Finalmente, la capa de servicios y datos se soporta sobre Supabase, utilizando "
        "sus componentes de autenticación, base de datos PostgreSQL y almacenamiento documental."
    )
    add_text(
        document,
        "Esta disposición por capas permite que la solución mantenga cohesión funcional y, al mismo tiempo, reduzca el acoplamiento entre "
        "interfaz, reglas de negocio y persistencia. De esta manera, la plataforma puede evolucionar sin comprometer su estabilidad general y "
        "facilita la incorporación de nuevos indicadores, nuevos formularios o nuevas rutas de análisis conforme avance el proyecto."
    )
    add_figure(
        document,
        16,
        "Arquitectura funcional de la plataforma web propuesta.",
        professional_arch,
        "Elaboración propia a partir de la solución desarrollada para la gestión del proyecto y de su integración operativa con Supabase.",
        width=6.8,
    )

    add_heading(document, "3.1.2. Esquema relacional de usuarios, roles y servicios de Supabase", 3)
    add_text(
        document,
        "La plataforma utiliza Supabase como infraestructura de soporte para la autenticación, la persistencia y la trazabilidad de eventos. "
        "En este entorno, la entidad de autenticación principal se encuentra en auth.users, mientras que la información de perfil ampliado se "
        "mantiene en tablas propias del esquema público, como perfiles_usuario, roles y roles_usuario. Esta separación permite diferenciar entre "
        "credenciales de acceso y atributos de negocio, garantizando una organización más normalizada y controlada de la información."
    )
    add_text(
        document,
        "La asignación de roles constituye una pieza central del modelo, ya que determina la visibilidad de páginas, acciones y permisos dentro "
        "del sistema. La administradora dispone de acceso integral al tablero, la configuración del proyecto, la gestión de usuarios, el diseño "
        "de cursos, la malla formativa y la producción científica. La formadora opera principalmente sobre los cursos, tareas, materiales y acompañamiento "
        "académico. La investigadora participa en la gestión analítica y científica del proyecto, mientras que la mujer emprendedora accede a su proceso "
        "formativo, a las entregas y al seguimiento de su participación. Este patrón relacional refuerza el principio de control por responsabilidades "
        "y favorece la consistencia de la seguridad a nivel de aplicación."
    )
    add_text(
        document,
        "Adicionalmente, Supabase cumple una función transversal al vincular tablas operativas como cuestionario_limpio_respuestas, cursos, modulos_curso, "
        "tareas_curso, entregas_tarea, documentos_proyecto, productos_cientificos, actividades_participante, notificaciones e historial_ingresos. "
        "En consecuencia, no se trata únicamente de un servicio de base de datos, sino de la columna vertebral de integración del sistema, tanto para el "
        "almacenamiento como para la sincronización de los módulos visibles en la plataforma."
    )
    if SCHEMA_IMG.exists():
        add_figure(
            document,
            17,
            "Esquema relacional asociado a perfiles de usuario, asignación de roles e integración con Supabase.",
            SCHEMA_IMG,
            "Captura del esquema relacional utilizado para conectar autenticación, perfiles y control de roles dentro de la plataforma.",
            width=6.3,
        )

    add_heading(document, "3.2. Módulos implementados dentro de la plataforma", 2)
    add_text(
        document,
        "La plataforma fue organizada en módulos funcionales para facilitar la operación del proyecto y mantener una correspondencia directa entre "
        "los procesos institucionales y las herramientas digitales implementadas. A continuación, se describen de forma individual los principales módulos desarrollados."
    )

    module_descriptions = [
        (
            "3.2.1. Módulo Dashboard",
            "El dashboard consolida los indicadores estratégicos del sistema y presenta de forma visual el avance del proyecto, la producción científica, "
            "la validación del programa, las próximas actividades y los gráficos de lectura diagnóstica. Su función es traducir los registros almacenados "
            "en Supabase en una vista ejecutiva de seguimiento para la toma de decisiones."
        ),
        (
            "3.2.2. Módulo Proyecto",
            "Este módulo permite organizar los documentos institucionales del proyecto, incluyendo carga, edición, eliminación y descarga de archivos. "
            "Su valor operativo reside en centralizar la documentación oficial en un repositorio accesible para los roles autorizados."
        ),
        (
            "3.2.3. Módulo Diagnóstico (Encuesta)",
            "El módulo de diagnóstico integra tanto la gestión analítica de resultados como la publicación del formulario de encuesta para captura de respuestas. "
            "Incluye lógica condicional, bloques temáticos, persistencia de respuestas y conexión directa con la base de datos diagnóstica."
        ),
        (
            "3.2.4. Módulo Analítica de Necesidades",
            "Transforma las respuestas de la encuesta en hallazgos visuales e interpretativos relacionados con necesidades de formación, competencias promedio, "
            "distribución territorial y tendencias de participación. Su objetivo es convertir la base de datos en evidencia útil para la planificación académica."
        ),
        (
            "3.2.5. Módulo Predicción de Cursos",
            "Procesa la información proveniente del cuestionario para sugerir cursos acordes con el perfil de las emprendedoras, sus sectores económicos, "
            "su modalidad preferida y las brechas más frecuentes detectadas. Este módulo aporta una capa de apoyo a decisiones curriculares."
        ),
        (
            "3.2.6. Módulo Diseño de Cursos",
            "Permite crear, editar, ocultar o eliminar cursos, así como estructurar módulos, contenidos y tareas. Constituye el núcleo de configuración pedagógica "
            "de la oferta formativa que luego se refleja en la experiencia de la mujer emprendedora."
        ),
        (
            "3.2.7. Módulo Malla Formativa",
            "Gestiona materiales académicos complementarios como sílabos, guías docentes, guías de estudio, rúbricas y fichas técnicas. Su propósito es respaldar "
            "la calidad documental y curricular de cada curso diseñado en la plataforma."
        ),
        (
            "3.2.8. Módulo Producción Científica",
            "Registra artículos, ponencias y demás productos académicos asociados al proyecto, junto con sus responsables, estados y fechas relevantes. "
            "Este módulo se enlaza con los indicadores del dashboard y con las próximas actividades del proyecto."
        ),
        (
            "3.2.9. Módulo Avance del Proyecto",
            "Administra actividades programadas y completadas, facilitando la visualización del progreso global, la identificación de pendientes y la relación "
            "temporal entre planificación y ejecución."
        ),
        (
            "3.2.10. Módulo Reportes",
            "Genera salidas documentales en formatos descargables, incluyendo exportaciones de encuestas e informes resumen. Su función es traducir la operación "
            "del sistema en productos académicos y administrativos útiles para seguimiento y presentación institucional."
        ),
        (
            "3.2.11. Módulo Configuración",
            "Permite administrar datos generales del proyecto, gestión de usuarios e historial de ingresos. Este módulo sostiene la gobernanza interna del sistema "
            "y conserva la trazabilidad de los cambios realizados por los distintos actores."
        ),
    ]
    for heading_text, body in module_descriptions:
        add_module_subsection(document, heading_text, body)

    dashboard_image = ASSETS_DIR / "dashboard-real.png"
    if dashboard_image.exists():
        add_figure(
            document,
            18,
            "Vista general del dashboard principal de la plataforma implementada.",
            dashboard_image,
            "Captura de la interfaz operativa donde se consolidan los indicadores principales del proyecto.",
            width=6.8,
        )

    add_heading(document, "3.3. Resultados analíticos integrados en la plataforma", 2)
    add_text(
        document,
        f"Al momento de la elaboración de este capítulo, la tabla cuestionario_limpio_respuestas registra {total} respuestas válidas. "
        f"Este volumen de información permite construir una base analítica consistente para comprender la distribución territorial, los sectores económicos "
        "predominantes, el interés de participación y la modalidad de capacitación preferida por las emprendedoras."
    )
    if top_parroquias:
        add_text(
            document,
            "En la distribución territorial destacan "
            + "; ".join(f"{title_case_first(name)} con {value} registros" for name, value in top_parroquias)
            + ". Esta concentración evidencia la utilidad del sistema para reconocer territorios con mayor participación y para focalizar futuras estrategias de convocatoria."
        )
    if top_sectores:
        add_text(
            document,
            "Por sector económico, sobresalen "
            + "; ".join(f"{name} ({value})" for name, value in top_sectores)
            + ". Este patrón confirma la necesidad de planificar contenidos de capacitación alineados con las actividades productivas realmente representadas en la base diagnóstica."
        )
    if modalidades:
        add_text(
            document,
            f"En cuanto a modalidad preferida, la distribución actual es la siguiente: {modalidades_text}. Este comportamiento orienta la toma de decisiones sobre el formato de los cursos y sustenta la pertinencia de incorporar experiencias virtuales e híbridas."
        )
    if intereses:
        add_text(
            document,
            f"Respecto al interés en participar en el programa, los registros muestran: {interes_text}. Este indicador no solo informa la disposición inicial de las participantes, sino que también ayuda a interpretar barreras de acceso, pertinencia temática o disponibilidad de tiempo."
        )

    figure_number = 19
    for title, filename, note, width in [
        (
            "Distribución de encuestas registradas por parroquia.",
            "grafico_parroquias.png",
            "Elaboración propia con base en los datos consolidados en Supabase.",
            6.5,
        ),
        (
            "Sectores económicos predominantes entre las emprendedoras encuestadas.",
            "grafico_sectores.png",
            "Elaboración propia a partir de la variable sector_economico registrada en la encuesta.",
            6.5,
        ),
        (
            "Modalidad preferida para recibir las capacitaciones.",
            "grafico_modalidad.png",
            "Elaboración propia con base en las respuestas válidas asociadas a modalidad_preferida.",
            6.2,
        ),
        (
            "Interés declarado en participar en el programa de formación.",
            "grafico_interes.png",
            "Elaboración propia a partir de la variable interes_programa almacenada en la base diagnóstica.",
            6.2,
        ),
    ]:
        image_path = ASSETS_DIR / filename
        if image_path.exists():
            add_figure(document, figure_number, title, image_path, note, width=width)
            figure_number += 1

    add_heading(document, "3.4. Aporte del sistema a la toma de decisiones", 2)
    add_text(
        document,
        "El valor principal de la plataforma no radica solamente en digitalizar procesos dispersos, sino en conectar la captura de información con decisiones concretas sobre formación, seguimiento y producción académica. "
        "En ese sentido, el sistema transforma datos operativos en insumos estratégicos para priorizar cursos, monitorear actividades, evaluar cobertura y justificar cambios metodológicos sobre la base de evidencia observable."
    )
    add_text(
        document,
        "La integración entre dashboard, analítica de necesidades, predicción de cursos, diseño curricular y avance del proyecto establece un circuito de retroalimentación continua. "
        "Cada nueva respuesta de la encuesta o actualización de una actividad influye en los indicadores, en las visualizaciones y en la lectura general del estado del proyecto, lo que fortalece el carácter dinámico y aplicado de la propuesta."
    )

    add_heading(document, "3.5. Conclusiones del capítulo", 2)
    add_text(
        document,
        "La propuesta tecnológica desarrollada demuestra que es factible consolidar en una sola plataforma los componentes operativos y analíticos requeridos por el proyecto Kawsay Emprende Guayaquil. "
        "La solución construida integra autenticación, roles, gestión documental, captura diagnóstica, análisis de necesidades, organización formativa, seguimiento de actividades y producción científica dentro de un mismo entorno digital."
    )
    add_text(
        document,
        "Asimismo, la articulación con Supabase aporta consistencia al modelo de datos y permite que la información capturada alimente de forma real los indicadores y gráficos visibles en la plataforma. "
        "Esto convierte al sistema en una herramienta no solo administrativa, sino también investigativa, al facilitar una lectura analítica permanente del comportamiento de la población objetivo."
    )

    add_heading(document, "3.6. Recomendaciones", 2)
    add_text(
        document,
        "Se recomienda mantener procesos periódicos de revisión y limpieza de datos, especialmente en variables territoriales y sectoriales, para sostener la calidad analítica de los resultados y mejorar la precisión de las decisiones formativas derivadas del sistema."
    )
    add_text(
        document,
        "También se recomienda ampliar el uso de la plataforma hacia mecanismos de seguimiento longitudinal que permitan relacionar el diagnóstico inicial con la permanencia en los cursos, los logros de aprendizaje y la evolución de los productos científicos generados durante la ejecución del proyecto."
    )

    document.save(OUTPUT_DOC)
    print(f"Documento generado: {OUTPUT_DOC}")

    with ZipFile(OUTPUT_DOC) as zf:
        media = [name for name in zf.namelist() if name.startswith("word/media/")]
        print(f"Imágenes embebidas: {len(media)}")


if __name__ == "__main__":
    main()
