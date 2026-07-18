from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


BASE_DIR = Path(r"D:\Proyecto_FCI_2025")
OUTPUT_DIR = BASE_DIR / "output" / "manual-uso-2026-07-16"
SCREENSHOTS_DIR = OUTPUT_DIR / "screenshots"
DOCX_PATH = OUTPUT_DIR / "Manual_de_uso_Plataforma_Kawsay_Emprende_Guayaquil.docx"
LOGO_PATH = BASE_DIR / "public" / "ug-192x192.png"

PRIMARY = RGBColor(13, 76, 140)
SECONDARY = RGBColor(60, 73, 89)
ACCENT = RGBColor(40, 178, 219)
LIGHT = RGBColor(233, 242, 250)
DARK = RGBColor(28, 38, 51)
WHITE = RGBColor(255, 255, 255)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_paragraph_border(paragraph, color: str = "D6E4F0", size: str = "6") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def set_font(run, size: float, color: RGBColor, bold: bool = False, italic: bool = False) -> None:
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    run.italic = italic


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = DARK

    for style_name, size, color in [
        ("Heading 1", 17, PRIMARY),
        ("Heading 2", 13, PRIMARY),
        ("Heading 3", 11.5, SECONDARY),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True


def add_header_footer(section) -> None:
    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run("Universidad de Guayaquil | Proyecto FCI 2025")
    set_font(run, 9, SECONDARY, bold=True)

    footer = section.footer
    table = footer.add_table(rows=1, cols=2, width=Inches(6.8))
    left = table.rows[0].cells[0].paragraphs[0]
    left.alignment = WD_ALIGN_PARAGRAPH.LEFT
    left_run = left.add_run("Manual de uso de la plataforma Kawsay Emprende Guayaquil")
    set_font(left_run, 8.5, SECONDARY)

    right = table.rows[0].cells[1].paragraphs[0]
    right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    page_run = right.add_run("Documento de apoyo operativo")
    set_font(page_run, 8.5, SECONDARY)


def add_cover(doc: Document) -> None:
    if LOGO_PATH.exists():
        doc.add_picture(str(LOGO_PATH), width=Inches(1.35))
        last = doc.paragraphs[-1]
        last.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Manual de uso de la plataforma")
    set_font(r, 18, SECONDARY, bold=False)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("Kawsay Emprende Guayaquil")
    set_font(r2, 24, PRIMARY, bold=True)

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run(
        "Sistema web para la gestión del proyecto de formación, análisis, cursos, "
        "producción científica y seguimiento operativo."
    )
    set_font(r3, 11.5, SECONDARY)

    info = doc.add_table(rows=4, cols=2)
    info.style = "Table Grid"
    rows = [
        ("Institución", "Universidad de Guayaquil"),
        ("Proyecto", "Proyecto FCI 2025"),
        ("Fecha del manual", "16 de julio de 2026"),
        ("Perfil revisado", "Administradora"),
    ]
    for idx, (label, value) in enumerate(rows):
        c1, c2 = info.rows[idx].cells
        c1.text = label
        c2.text = value
        set_cell_shading(c1, "E9F2FA")
        for paragraph in c1.paragraphs + c2.paragraphs:
            for run in paragraph.runs:
                set_font(run, 10.5, DARK, bold=(paragraph in c1.paragraphs))

    doc.add_paragraph("")
    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    nr = note.add_run(
        "Este manual describe el uso funcional de la plataforma a partir de una revisión "
        "real del sistema publicado en localhost con navegación y capturas operativas."
    )
    set_font(nr, 9.5, SECONDARY, italic=True)
    doc.add_page_break()


def add_title(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(text)
    set_font(r, 19, PRIMARY, bold=True)
    p.paragraph_format.space_after = Pt(3)
    set_paragraph_border(p, color="6AAED6", size="8")


def add_body(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    r = p.add_run(text)
    set_font(r, 10.5, DARK)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(item)
        set_font(r, 10.5, DARK)


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(item)
        set_font(r, 10.5, DARK)


def add_callout(doc: Document, title: str, text: str, fill: str = "E9F2FA") -> None:
    table = doc.add_table(rows=1, cols=1)
    table.autofit = True
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, fill)
    p1 = cell.paragraphs[0]
    rr = p1.add_run(title + "\n")
    set_font(rr, 10.5, PRIMARY, bold=True)
    rr2 = p1.add_run(text)
    set_font(rr2, 10, DARK)


def add_screenshot(doc: Document, path: Path, caption: str) -> None:
    if not path.exists():
        return
    doc.add_picture(str(path), width=Inches(6.3))
    image = doc.paragraphs[-1]
    image.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp = doc.add_paragraph()
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cr = cp.add_run(caption)
    set_font(cr, 9.5, SECONDARY, italic=True)
    cp.paragraph_format.space_after = Pt(8)


def add_module_section(doc: Document, title: str, purpose: str, actions: list[str], image_name: str, caption: str) -> None:
    doc.add_paragraph(title, style="Heading 1")
    add_body(doc, purpose)
    add_bullets(doc, actions)
    add_screenshot(doc, SCREENSHOTS_DIR / image_name, caption)


def build_manual() -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    style_document(doc)
    add_header_footer(doc.sections[0])
    add_cover(doc)

    add_title(doc, "1. Objetivo del manual")
    add_body(
        doc,
        "Este manual explica el funcionamiento operativo de la plataforma Kawsay Emprende "
        "Guayaquil desde la perspectiva de administración del proyecto. Su propósito es "
        "facilitar el uso diario del sistema, reducir errores de operación y dejar documentado "
        "cómo se navega, qué información se registra y cómo se relacionan los módulos."
    )
    add_callout(
        doc,
        "Alcance del documento",
        "La guía cubre acceso, panel principal, documentos del proyecto, análisis de encuestas, "
        "predicción de cursos, diseño académico, producción científica, avance del proyecto, "
        "reportes y configuración.",
    )

    doc.add_paragraph("2. Perfil y alcance funcional", style="Heading 1")
    add_body(
        doc,
        "La revisión se realizó con un usuario autenticado como Administradora. Este perfil "
        "dispone de acceso a las secciones estratégicas de la plataforma y puede registrar, "
        "editar, visualizar y exportar información clave del proyecto."
    )
    add_bullets(
        doc,
        [
            "Administradora: controla la configuración general, usuarios, documentos, cursos, reportes y seguimiento.",
            "Formadora e investigadora: participan en módulos específicos como cursos, producción y evaluación.",
            "Mujer emprendedora: interactúa principalmente con formación, tareas, notificaciones y encuesta publicada.",
            "Encuesta pública: puede responderse sin iniciar sesión desde la ruta publicada del cuestionario.",
        ],
    )

    doc.add_paragraph("3. Estructura general de navegación", style="Heading 1")
    add_body(
        doc,
        "La aplicación organiza sus funciones en un panel lateral izquierdo y un encabezado superior. "
        "El panel concentra el acceso a las secciones operativas, mientras que el encabezado muestra "
        "la identidad del proyecto, notificaciones y acceso al perfil."
    )
    add_bullets(
        doc,
        [
            "Panel lateral: permite cambiar entre módulos sin perder el contexto general del proyecto.",
            "Encabezado superior: muestra el nombre del proyecto, descripción, campana de notificaciones y menú de perfil.",
            "Contenido central: despliega formularios, gráficos, tablas, cronogramas y acciones del módulo seleccionado.",
        ],
    )
    add_screenshot(
        doc,
        SCREENSHOTS_DIR / "01-dashboard.png",
        "Figura 1. Vista general del panel principal con barra lateral, encabezado y tablero ejecutivo.",
    )

    doc.add_paragraph("4. Flujo de uso recomendado", style="Heading 1")
    add_numbered(
        doc,
        [
            "Acceder a la plataforma con las credenciales institucionales.",
            "Revisar el Dashboard para identificar avance, producción, validación y actividades próximas.",
            "Actualizar la documentación y la información del proyecto cuando exista nueva evidencia.",
            "Monitorear el diagnóstico y las métricas relevantes conforme se registran nuevas encuestas.",
            "Analizar la predicción de cursos y trasladar decisiones al módulo de diseño de cursos.",
            "Registrar producción científica y actividades para mantener sincronizados los indicadores de seguimiento.",
            "Generar reportes cuando se requiera sustento documental o resumen ejecutivo.",
            "Usar Configuración para ajuste de usuarios, historial de ingresos y datos generales del proyecto.",
        ],
    )

    add_module_section(
        doc,
        "5. Dashboard",
        "El Dashboard resume el estado operativo del proyecto. Desde esta pantalla se visualizan "
        "indicadores de avance, producción científica, validación del programa, necesidades de formación, "
        "competencias promedio y actividades próximas. Es la vista inicial para control y toma de decisiones.",
        [
            "Tarjeta Avance del Proyecto: muestra porcentaje global, estado actual y fechas de inicio y fin.",
            "Tarjeta Cursos Diseñados: refleja el número total de cursos creados y su estado general.",
            "Tarjeta Producción Científica: concentra ejecutados, totales, porcentaje de cumplimiento, en proceso y pendientes.",
            "Tarjeta Validación del Programa: contabiliza participantes encuestadas y compara el resultado con la meta definida.",
            "Gráfico Avance del Proyecto en el Tiempo: compara planificado frente a ejecutado.",
            "Gráfico Producción por Investigador: distribuye productos ejecutados según tipo y responsable.",
            "Radar de Competencias y gráfico de necesidades: sintetizan la información derivada del cuestionario.",
            "Próximas Actividades: integra registros programados y en ejecución provenientes de producción científica y avance del proyecto.",
        ],
        "01-dashboard.png",
        "Figura 2. Dashboard ejecutivo con indicadores, gráficos y actividades del proyecto.",
    )

    add_module_section(
        doc,
        "6. Proyecto",
        "El módulo Proyecto centraliza la documentación institucional. Permite registrar, actualizar y consultar "
        "archivos de soporte como resoluciones, cronogramas, informes parciales o documentos metodológicos.",
        [
            "Subir documentos del proyecto con nombre descriptivo y clasificación funcional.",
            "Editar registros documentales cuando se deba reemplazar un archivo o ajustar sus datos.",
            "Eliminar documentos obsoletos o cargados por error.",
            "Descargar el archivo conservando su nombre original para revisión o respaldo.",
        ],
        "02-proyecto.png",
        "Figura 3. Gestión documental del proyecto con carga, edición y descarga de archivos.",
    )

    add_module_section(
        doc,
        "7. Diagnóstico (Encuesta)",
        "Esta sección presenta el tablero analítico del cuestionario diagnóstico. Se orienta al seguimiento de "
        "respuestas por bloque, distribución de categorías y lectura de resultados relevantes para el proyecto.",
        [
            "Visualizar resultados por bloque del instrumento aplicado a mujeres emprendedoras.",
            "Revisar gráficos de distribución de respuestas y tendencias principales.",
            "Utilizar segmentaciones y comparaciones para interpretar datos del cuestionario.",
            "Acceder al formulario publicado para su respuesta desde una ruta pública independiente.",
        ],
        "03-diagnostico-dashboard.png",
        "Figura 4. Panel de diagnóstico con resultados y distribución de respuestas de la encuesta.",
    )

    doc.add_paragraph("8. Encuesta pública publicada", style="Heading 1")
    add_body(
        doc,
        "El cuestionario diagnóstico también se encuentra disponible como formulario público. Se usa para "
        "captura directa de respuestas sin necesidad de autenticación y envía la información a la base de datos "
        "del proyecto para análisis posterior."
    )
    add_bullets(
        doc,
        [
            "Organiza las preguntas por bloques temáticos con navegación secuencial.",
            "Aplica condiciones lógicas para mostrar u ocultar campos según las respuestas.",
            "Valida que las preguntas obligatorias estén completas antes de avanzar.",
            "Al finalizar, presenta confirmación de envío y evita modificaciones posteriores desde el mismo dispositivo.",
        ],
    )
    add_screenshot(
        doc,
        SCREENSHOTS_DIR / "11-encuesta-publica.png",
        "Figura 5. Formulario público del cuestionario diagnóstico por bloques temáticos.",
    )

    add_module_section(
        doc,
        "9. Métricas relevantes",
        "Métricas relevantes consolida la lectura analítica derivada de las encuestas. Permite identificar "
        "necesidades formativas críticas, tendencias de competencias y relaciones útiles para la toma de decisiones.",
        [
            "Identificar la necesidad principal y su peso relativo dentro del conjunto de respuestas.",
            "Comparar bloques de competencias para detectar fortalezas y brechas.",
            "Analizar perfiles agregados sin depender únicamente del Dashboard.",
            "Servir como base interpretativa para la selección y diseño de cursos formativos.",
        ],
        "04-metricas-relevantes.png",
        "Figura 6. Módulo de métricas relevantes con lectura analítica de necesidades y competencias.",
    )

    add_module_section(
        doc,
        "10. Predicción de Cursos",
        "Este módulo propone cursos a partir del análisis de los datos del cuestionario. Su objetivo es "
        "transformar la evidencia levantada en sugerencias de formación más pertinentes para las emprendedoras.",
        [
            "Procesa variables como parroquia, sector económico, ingreso mensual y nivel educativo.",
            "Presenta cursos sugeridos según necesidades detectadas y perfiles dominantes.",
            "Puede operar en modo rápido local o apoyarse en Gemini cuando la integración esté activada.",
            "Sirve como insumo para decidir qué cursos deben formalizarse en Diseño de Cursos.",
        ],
        "05-prediccion-cursos.png",
        "Figura 7. Predicción de cursos basada en datos del diagnóstico y sugerencias formativas.",
    )
    add_callout(
        doc,
        "Lectura funcional",
        "Cuando Gemini no está disponible o se alcanza el límite gratuito, la plataforma mantiene un modo "
        "de respaldo con predicción local para no interrumpir el análisis.",
        fill="F3F7FB",
    )

    add_module_section(
        doc,
        "11. Diseño de Cursos",
        "Diseño de Cursos permite crear la oferta formativa del proyecto. Aquí se definen título, descripción, "
        "estado, visibilidad, contenido formativo, tareas y materiales vinculados a cada curso.",
        [
            "Crear, editar, ocultar o eliminar cursos según la planificación académica.",
            "Construir el contenido interno de cada curso por módulos o bloques.",
            "Asignar tareas y recursos visibles para las mujeres emprendedoras.",
            "Controlar el estado del curso para su posterior publicación o validación.",
        ],
        "06-diseno-cursos.png",
        "Figura 8. Constructor de cursos con edición de contenidos, materiales y tareas.",
    )

    add_module_section(
        doc,
        "12. Producción Científica",
        "El módulo registra productos científicos del proyecto, como artículos de alto impacto o regionales. "
        "Su información alimenta el Dashboard y también se refleja en las actividades programadas.",
        [
            "Registrar título del producto, tipo, responsables y estado del proceso.",
            "Diferenciar productos pendientes, en redacción, en revisión y ejecutados.",
            "Guardar fechas objetivo y fecha efectiva de publicación para seguimiento.",
            "Adjuntar evidencia o enlaces asociados cuando corresponda.",
        ],
        "07-produccion-cientifica.png",
        "Figura 9. Registro y seguimiento de productos científicos del proyecto.",
    )

    add_module_section(
        doc,
        "13. Avance del Proyecto",
        "Avance del Proyecto se concentra en la planificación operativa. Permite crear actividades, editar su "
        "estado y consolidar un seguimiento temporal que se integra con el panel principal.",
        [
            "Crear actividades del proyecto con título, descripción, orden, fecha objetivo y estado.",
            "Editar actividades existentes para actualizar cumplimiento o reprogramación.",
            "Mantener sincronizada la línea de tiempo del Dashboard con las actividades creadas.",
            "Consultar actividades programadas y distinguir las completadas de las aún no ejecutadas.",
        ],
        "08-avance-proyecto.png",
        "Figura 10. Gestión del avance operativo con actividades programadas y edición de estados.",
    )

    add_module_section(
        doc,
        "14. Reportes",
        "Reportes concentra las salidas descargables del sistema. Desde aquí se exporta la base de encuestas y "
        "también se genera un informe ejecutivo en PDF con base en la información consolidada del proyecto.",
        [
            "Descargar la encuesta completa en formato Excel con encabezados descriptivos.",
            "Emitir un PDF resumen con información del Dashboard y lectura ejecutiva del estado del proyecto.",
            "Usar los reportes como soporte para reuniones, revisión técnica y avance institucional.",
        ],
        "09-reportes.png",
        "Figura 11. Página de reportes con exportación de encuestas y generación de informes ejecutivos.",
    )

    add_module_section(
        doc,
        "15. Configuración",
        "Configuración agrupa las opciones administrativas del sistema. Permite gestionar usuarios, revisar "
        "historial de ingresos y mantener actualizados los datos generales del proyecto.",
        [
            "Gestión de usuarios: crear, invitar, editar nombre y asignar rol institucional.",
            "Historial de ingresos: revisar accesos, navegación y actividad registrada por rol.",
            "Proyecto: ajustar nombre, descripción, fechas y meta de validación cuando corresponda.",
        ],
        "10-configuracion.png",
        "Figura 12. Configuración administrativa del sistema con usuarios, historial y parámetros generales.",
    )

    doc.add_paragraph("16. Integración entre módulos", style="Heading 1")
    add_body(
        doc,
        "La plataforma no funciona como un conjunto de páginas aisladas. Varios módulos alimentan de forma "
        "directa los indicadores del tablero principal y los reportes ejecutivos."
    )
    add_bullets(
        doc,
        [
            "Las respuestas del cuestionario alimentan diagnóstico, métricas relevantes, validación y predicción de cursos.",
            "La predicción de cursos apoya la decisión que luego se concreta en Diseño de Cursos.",
            "Producción Científica y Avance del Proyecto sincronizan actividades visibles en el Dashboard.",
            "Configuración del proyecto define nombre, descripción, fechas y meta usados por el encabezado y tarjetas.",
            "Reportes consolida la información de los demás módulos en formatos descargables.",
        ],
    )

    doc.add_paragraph("17. Recomendaciones de uso", style="Heading 1")
    add_bullets(
        doc,
        [
            "Actualizar primero la información base del proyecto antes de revisar indicadores.",
            "Revisar periódicamente el Dashboard para detectar atrasos o metas cercanas.",
            "Mantener consistencia entre Producción Científica y Avance del Proyecto para evitar duplicidades.",
            "Cargar archivos con nombres claros y definitivos en el módulo Proyecto.",
            "Usar Reportes como salida oficial para reuniones y seguimiento institucional.",
        ],
    )

    doc.add_paragraph("18. Cierre", style="Heading 1")
    add_body(
        doc,
        "La plataforma Kawsay Emprende Guayaquil integra en un mismo entorno el seguimiento técnico, formativo "
        "y analítico del proyecto FCI 2025. Su valor operativo radica en que concentra evidencias, resultados, "
        "cursos, cronogramas y reportes bajo una lógica conectada, lo cual facilita el control administrativo y "
        "la toma de decisiones informadas."
    )

    doc.save(str(DOCX_PATH))


if __name__ == "__main__":
    build_manual()
